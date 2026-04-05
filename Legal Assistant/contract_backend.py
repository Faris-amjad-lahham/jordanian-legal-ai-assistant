import os
import re
from dataclasses import dataclass, field
from typing import Dict, List, Optional, Any, Tuple
from datetime import datetime, date

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Optional OpenAI
try:
    from openai import OpenAI
    from openai.types.chat import ChatCompletionMessageParam
    _openai_client = OpenAI()
except Exception:
    _openai_client = None  # type: ignore
    from typing import TYPE_CHECKING
    if TYPE_CHECKING:
        from openai.types.chat import ChatCompletionMessageParam  # type: ignore

EXPORT_DIR = "exports"
os.makedirs(EXPORT_DIR, exist_ok=True)

MAX_VERSIONS = 10  # Increased from 5

# =============================================
# JORDANIAN LABOUR LAW CONSTANTS (for validation)
# =============================================
JO_MIN_WAGE_JOD = 260          # Minimum wage as of 2023 (JOD/month)
JO_MAX_PROBATION_DAYS = 90     # Article 25: max 3 months probation
JO_MAX_DAILY_HOURS = 8         # Article 56: normal working hours
JO_MAX_WEEKLY_HOURS = 48       # Article 56: max weekly hours
JO_ANNUAL_LEAVE_DAYS = 14      # Article 61: min annual leave (after 1 year)
JO_NOTICE_PERIOD_DAYS = 30     # Article 32: min notice period (1+ year service)
JO_MAX_FIXED_CONTRACT_YEARS = 2  # Typical max for fixed-term contracts

DISCLAIMER_EN = (
    "LEGAL DISCLAIMER:\n"
    "This document is a NON-BINDING employment contract draft for informational purposes only.\n"
    "It does not constitute legal advice. Consult a licensed lawyer in Jordan before signing any contract.\n"
    "This draft has been generated automatically and may not reflect all applicable Jordanian laws.\n"
)

DISCLAIMER_AR = (
    "تنويه قانوني:\n"
    "هذه الوثيقة مسودة عقد عمل غير ملزمة قانونياً وللأغراض المعلوماتية فقط.\n"
    "لا تُعدّ استشارة قانونية. يُنصح بمراجعة محامٍ مرخص في الأردن قبل توقيع أي عقد.\n"
    "تم إنشاء هذه المسودة تلقائياً وقد لا تعكس جميع قوانين العمل الأردنية المعمول بها.\n"
)

# =============================================
# HELPERS
# =============================================
def _now_iso() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


def _ensure_disclaimer(text: str, lang: str = "en") -> str:
    if not text:
        return text
    disclaimer = DISCLAIMER_AR if lang == "ar" else DISCLAIMER_EN
    if "LEGAL DISCLAIMER:" in text or "تنويه قانوني:" in text:
        return text
    return text.rstrip() + "\n\n" + disclaimer


def _is_arabic(text: str) -> bool:
    arabic_chars = sum(1 for c in (text or "") if "\u0600" <= c <= "\u06FF")
    total = len([c for c in (text or "") if c.strip()])
    return (arabic_chars / total) > 0.3 if total else False


def _parse_date(date_str: str) -> Optional[date]:
    """Try to parse a date string in common formats."""
    formats = ["%d/%m/%Y", "%Y-%m-%d", "%d-%m-%Y", "%m/%d/%Y", "%d.%m.%Y"]
    for fmt in formats:
        try:
            return datetime.strptime(date_str.strip(), fmt).date()
        except ValueError:
            continue
    return None


# =============================================
# DATA CLASSES
# =============================================
@dataclass
class ValidationIssue:
    severity: str       # "error" | "warning" | "info"
    field: str
    message: str
    article_ref: Optional[str] = None


@dataclass
class ContractVersion:
    text: str
    created_at: str
    data: Dict[str, Any] = field(default_factory=dict)
    meta: Dict[str, Any] = field(default_factory=dict)


@dataclass
class ContractStore:
    current: Optional[str] = None
    current_data: Dict[str, Any] = field(default_factory=dict)
    versions: List[ContractVersion] = field(default_factory=list)
    chat_history: List[Dict[str, str]] = field(default_factory=list)
    lang: str = "en"

    def save_version(self, meta: Optional[Dict[str, Any]] = None) -> None:
        if self.current:
            self.versions.insert(0, ContractVersion(
                text=self.current,
                created_at=_now_iso(),
                data=dict(self.current_data),
                meta=meta or {}
            ))
            self.versions = self.versions[:MAX_VERSIONS]

    def get_previous(self, index: int = 0) -> Optional[str]:
        if 0 <= index < len(self.versions):
            return self.versions[index].text
        return None

    def restore_version(self, index: int) -> bool:
        if 0 <= index < len(self.versions):
            self.save_version(meta={"action": "restore", "from_index": index})
            self.current = self.versions[index].text
            self.current_data = dict(self.versions[index].data)
            return True
        return False

    def add_chat_message(self, role: str, content: str) -> None:
        self.chat_history.append({"role": role, "content": content})
        # Keep last 20 messages
        if len(self.chat_history) > 20:
            self.chat_history = self.chat_history[-20:]

    def completeness_score(self) -> int:
        """Return % of required fields filled."""
        required = ["employer", "employee", "job_title", "location",
                    "contract_type", "start_date", "salary", "hours", "days"]
        filled = sum(1 for f in required if self.current_data.get(f, "").strip())
        return int((filled / len(required)) * 100)


# =============================================
# VALIDATION ENGINE
# =============================================
def validate_contract_data(data: Dict[str, str]) -> List[ValidationIssue]:
    """
    Validate contract fields against Jordanian Labour Law.
    Returns list of ValidationIssue (errors, warnings, info).
    """
    issues: List[ValidationIssue] = []

    # --- Required fields ---
    required_fields = {
        "employer": "Employer name",
        "employee": "Employee name",
        "job_title": "Job title",
        "location": "Work location",
        "contract_type": "Contract type",
        "start_date": "Start date",
        "salary": "Salary",
        "hours": "Working hours per day",
        "days": "Working days per week",
    }
    for field_key, label in required_fields.items():
        if not data.get(field_key, "").strip():
            issues.append(ValidationIssue(
                severity="error",
                field=field_key,
                message=f"{label} is required.",
            ))

    # --- Social Security reminder (always required in Jordan) ---
    issues.append(ValidationIssue(
        severity="info",
        field="social_security",
        message="Both parties are legally required to register with Jordanian Social Security (SSC). "
                "Ensure SSC deductions are included in the contract.",
        article_ref="Social Security Law No. 1 / 2014"
    ))

    # --- Salary validation (Article 52) ---
    salary_str = data.get("salary", "").strip()
    if salary_str:
        try:
            salary = float(re.sub(r"[^\d.]", "", salary_str))
            if salary < JO_MIN_WAGE_JOD:
                issues.append(ValidationIssue(
                    severity="error",
                    field="salary",
                    message=f"Salary {salary:.0f} JOD is below the Jordanian minimum wage of {JO_MIN_WAGE_JOD} JOD/month.",
                    article_ref="Article 52 — Minimum Wage"
                ))
            elif salary < JO_MIN_WAGE_JOD * 1.2:
                issues.append(ValidationIssue(
                    severity="warning",
                    field="salary",
                    message=f"Salary is close to the minimum wage ({JO_MIN_WAGE_JOD} JOD). Verify this is appropriate for the role.",
                    article_ref="Article 52 — Minimum Wage"
                ))
        except (ValueError, TypeError):
            issues.append(ValidationIssue(
                severity="error",
                field="salary",
                message="Salary must be a valid number (e.g., 500).",
            ))

    # --- Working hours (Article 56) ---
    hours_str = data.get("hours", "").strip()
    if hours_str:
        try:
            hours = float(re.sub(r"[^\d.]", "", hours_str))
            if hours > JO_MAX_DAILY_HOURS:
                issues.append(ValidationIssue(
                    severity="warning",
                    field="hours",
                    message=f"{hours:.0f} hours/day exceeds the standard {JO_MAX_DAILY_HOURS} hours. Overtime rules will apply.",
                    article_ref="Article 56 — Working Hours"
                ))
            if hours <= 0:
                issues.append(ValidationIssue(
                    severity="error",
                    field="hours",
                    message="Working hours must be greater than 0.",
                ))
        except (ValueError, TypeError):
            issues.append(ValidationIssue(
                severity="error",
                field="hours",
                message="Working hours must be a valid number.",
            ))

    # --- Working days + weekly hours check (Article 56) ---
    days_str = data.get("days", "").strip()
    if hours_str and days_str:
        try:
            hours = float(re.sub(r"[^\d.]", "", hours_str))
            days = float(re.sub(r"[^\d.]", "", days_str))
            weekly_hours = hours * days
            if weekly_hours > JO_MAX_WEEKLY_HOURS:
                issues.append(ValidationIssue(
                    severity="error",
                    field="days",
                    message=f"Total weekly hours ({weekly_hours:.0f}h) exceeds the legal maximum of {JO_MAX_WEEKLY_HOURS} hours.",
                    article_ref="Article 56 — Maximum Weekly Hours"
                ))
            if days > 7:
                issues.append(ValidationIssue(
                    severity="error",
                    field="days",
                    message="Working days cannot exceed 7 days per week.",
                ))
            if days == 7:
                issues.append(ValidationIssue(
                    severity="warning",
                    field="days",
                    message="Working 7 days/week is unusual. Employees are entitled to at least one weekly rest day.",
                    article_ref="Article 57 — Weekly Rest"
                ))
        except (ValueError, TypeError):
            pass

    # --- Date validation ---
    start_str = data.get("start_date", "").strip()
    end_str = data.get("end_date", "").strip()
    contract_type = data.get("contract_type", "").strip().lower()

    start_date_obj = _parse_date(start_str) if start_str else None
    end_date_obj = _parse_date(end_str) if end_str else None

    if start_str and not start_date_obj:
        issues.append(ValidationIssue(
            severity="error",
            field="start_date",
            message="Start date format not recognized. Use DD/MM/YYYY (e.g., 01/06/2026).",
        ))

    if "limited" in contract_type:
        if not end_str:
            issues.append(ValidationIssue(
                severity="error",
                field="end_date",
                message="End date is required for Limited (fixed-term) contracts.",
                article_ref="Article 25 — Fixed-Term Contracts"
            ))
        elif end_date_obj and start_date_obj:
            if end_date_obj <= start_date_obj:
                issues.append(ValidationIssue(
                    severity="error",
                    field="end_date",
                    message="End date must be after start date.",
                ))
            else:
                duration_years = (end_date_obj - start_date_obj).days / 365
                if duration_years > JO_MAX_FIXED_CONTRACT_YEARS:
                    issues.append(ValidationIssue(
                        severity="warning",
                        field="end_date",
                        message=f"Contract duration ({duration_years:.1f} years) is long for a fixed-term contract. Consider whether unlimited contract is more appropriate.",
                        article_ref="Article 25 — Contract Duration"
                    ))

    # --- Probation period (Article 25) ---
    probation_str = data.get("probation_days", "").strip()
    if probation_str:
        try:
            prob_days = int(re.sub(r"[^\d]", "", probation_str))
            if prob_days > JO_MAX_PROBATION_DAYS:
                issues.append(ValidationIssue(
                    severity="error",
                    field="probation_days",
                    message=f"Probation period ({prob_days} days) exceeds the legal maximum of {JO_MAX_PROBATION_DAYS} days (3 months).",
                    article_ref="Article 25 — Probation Period"
                ))
        except (ValueError, TypeError):
            pass

    # --- Termination notice (Article 32) ---
    termination_str = data.get("termination", "").strip()
    if termination_str:
        notice_match = re.search(r"(\d+)\s*day", termination_str.lower())
        if notice_match:
            notice_days = int(notice_match.group(1))
            if notice_days < JO_NOTICE_PERIOD_DAYS:
                issues.append(ValidationIssue(
                    severity="warning",
                    field="termination",
                    message=f"Notice period ({notice_days} days) is less than the recommended {JO_NOTICE_PERIOD_DAYS} days under Jordanian law.",
                    article_ref="Article 32 — Notice Period"
                ))

    # --- Annual leave info (Article 61) ---
    issues.append(ValidationIssue(
        severity="info",
        field="annual_leave",
        message="Employees are entitled to 14 days paid annual leave/year (rising to 21 days after 5 continuous years). "
                "This must be reflected in the contract.",
        article_ref="Article 61 — Annual Leave"
    ))

    # --- Ramadan hours (Article 59) ---
    hours_str2 = data.get("hours", "").strip()
    if hours_str2:
        try:
            h = float(re.sub(r"[^\d.]", "", hours_str2))
            if h > 6:
                issues.append(ValidationIssue(
                    severity="warning",
                    field="hours",
                    message=f"During Ramadan, working hours must be reduced to 6 hours/day for Muslim employees "
                            f"(your contract specifies {h:.0f} hours). Consider adding a Ramadan clause.",
                    article_ref="Article 59 — Ramadan Working Hours"
                ))
        except (ValueError, TypeError):
            pass

    return issues


def get_validation_summary(issues: List[ValidationIssue]) -> Dict[str, Any]:
    errors = [i for i in issues if i.severity == "error"]
    warnings = [i for i in issues if i.severity == "warning"]
    infos = [i for i in issues if i.severity == "info"]
    return {
        "valid": len(errors) == 0,
        "error_count": len(errors),
        "warning_count": len(warnings),
        "info_count": len(infos),
        "errors": errors,
        "warnings": warnings,
        "infos": infos,
        "all_issues": issues,
    }


# =============================================
# CONTRACT GENERATION (AI-POWERED)
# =============================================
def generate_contract_ai(
    store: ContractStore,
    data: Dict[str, str],
    lang: str = "en",
) -> Tuple[str, List[ValidationIssue]]:
    """
    Generate contract using GPT with Labour Law validation.
    Returns (contract_text, validation_issues).
    """
    # Run validation first
    issues = validate_contract_data(data)
    validation_summary = get_validation_summary(issues)

    # Detect language
    detected_lang = "ar" if _is_arabic(data.get("employee", "") + data.get("employer", "")) else lang
    store.lang = detected_lang

    # Build validation context for GPT
    law_warnings = []
    for issue in issues:
        if issue.article_ref:
            law_warnings.append(f"- {issue.message} ({issue.article_ref})")

    law_context = "\n".join(law_warnings) if law_warnings else "No legal issues detected."

    if detected_lang == "ar":
        system_prompt = """أنت مساعد متخصص في صياغة عقود العمل الأردنية.
قم بصياغة عقد عمل غير ملزم قانونياً بناءً على البيانات المقدمة.
القواعد:
- استخدم اللغة العربية الرسمية
- اذكر أرقام المواد القانونية ذات الصلة
- أضف تنويهاً قانونياً في النهاية
- إذا كانت هناك مخاوف قانونية، أشر إليها بوضوح في العقد
- اجعل العقد منظماً ومقسماً بوضوح"""
    else:
        system_prompt = """You are a specialist in drafting Jordanian employment contracts.
Draft a non-binding employment contract based on the provided data.
Rules:
- Use formal, professional English
- Reference relevant Jordanian Labour Law article numbers
- Add a legal disclaimer at the end
- If there are legal concerns, note them clearly in the contract
- Structure the contract with clear numbered sections"""

    user_prompt = f"""
Contract Details:
{_format_data_for_prompt(data)}

Legal Considerations (from Jordanian Labour Law):
{law_context}

Please draft a complete, professional employment contract. Include all standard sections:
1. Parties
2. Position & Duties  
3. Contract Term
4. Compensation & Benefits
5. Working Hours
6. Leave Entitlements (per Jordanian Labour Law)
7. Termination Conditions
8. Governing Law (Jordanian Labour Law No. 8 of 1996)
9. Legal Disclaimer
"""

    # Save previous version
    store.save_version(meta={"action": "generate"})
    store.current_data = dict(data)
    store.add_chat_message("user", f"Generate contract for {data.get('employee', 'employee')}")

    if _openai_client is None or not os.getenv("OPENAI_API_KEY"):
        # Fallback: rule-based contract
        contract = _generate_contract_fallback(data, detected_lang)
        store.current = _ensure_disclaimer(contract, detected_lang)
        store.add_chat_message("assistant", "Contract generated (offline mode - no API key).")
        return store.current, issues

    messages = [  # type: ignore
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt},
    ]

    resp = _openai_client.chat.completions.create(
        model="gpt-4o-mini",
        messages=messages,  # type: ignore
        temperature=0,
    )
    contract = (resp.choices[0].message.content or "").strip()
    contract = _ensure_disclaimer(contract, detected_lang)

    store.current = contract
    store.add_chat_message("assistant", "Contract generated successfully.")

    return store.current, issues


def _format_data_for_prompt(data: Dict[str, str]) -> str:
    labels = {
        "employer": "Employer",
        "employee": "Employee",
        "job_title": "Job Title",
        "location": "Work Location",
        "contract_type": "Contract Type",
        "start_date": "Start Date",
        "end_date": "End Date",
        "salary": "Monthly Salary (JOD)",
        "hours": "Working Hours/Day",
        "days": "Working Days/Week",
        "probation_days": "Probation Period (Days)",
        "termination": "Termination Terms",
        "benefits": "Additional Benefits",
    }
    lines = []
    for key, label in labels.items():
        val = data.get(key, "").strip()
        if val:
            lines.append(f"- {label}: {val}")
    return "\n".join(lines)


def _generate_contract_fallback(data: Dict[str, str], lang: str = "en") -> str:
    """Rule-based contract when no API key is available."""
    disclaimer = DISCLAIMER_AR if lang == "ar" else DISCLAIMER_EN
    return f"""NON-BINDING EMPLOYMENT CONTRACT DRAFT

This Employment Contract is made on: {datetime.now().strftime('%d/%m/%Y')}

PARTIES:
Employer: {data.get('employer', '___')}
Employee: {data.get('employee', '___')}

POSITION:
Job Title: {data.get('job_title', '___')}
Work Location: {data.get('location', '___')}

CONTRACT TERM:
Type: {data.get('contract_type', '___')}
Start Date: {data.get('start_date', '___')}
End Date: {data.get('end_date', 'N/A (Unlimited)')}

COMPENSATION:
Monthly Salary: {data.get('salary', '___')} JOD

WORKING HOURS:
Hours per Day: {data.get('hours', '___')}
Days per Week: {data.get('days', '___')}

LEAVE ENTITLEMENTS:
Annual Leave: 14 days/year (as per Article 61, Jordanian Labour Law)
Sick Leave: As per Article 65, Jordanian Labour Law

TERMINATION:
{data.get('termination', '30 days written notice required by either party.')}

GOVERNING LAW:
This contract is governed by Jordanian Labour Law No. 8 of 1996 and its amendments.

{disclaimer}""".strip()


# =============================================
# CONVERSATIONAL CONTRACT EDITOR (AI-POWERED)
# =============================================
def chat_edit_contract(
    store: ContractStore,
    user_message: str,
) -> Tuple[str, str, List[ValidationIssue]]:
    """
    Edit contract via natural language chat.
    Returns (updated_contract, assistant_reply, validation_issues).
    """
    if not store.current:
        reply = "No contract exists yet. Please generate a contract first."
        store.add_chat_message("assistant", reply)
        return "", reply, []

    store.add_chat_message("user", user_message)

    if _openai_client is None or not os.getenv("OPENAI_API_KEY"):
        reply = "⚠️ API key required for conversational editing."
        store.add_chat_message("assistant", reply)
        return store.current, reply, []

    lang = store.lang
    if lang == "ar":
        system_prompt = """أنت مساعد لتعديل عقود العمل الأردنية.
لديك العقد الحالي وطلب التعديل من المستخدم.
قم بـ:
1. تطبيق التعديل المطلوب على العقد
2. الحفاظ على هيكل العقد وتنسيقه
3. التأكد من التوافق مع قانون العمل الأردني
4. الإبقاء على التنويه القانوني في النهاية
5. الرد بالعقد المعدّل كاملاً فقط، بدون تعليقات إضافية"""
    else:
        system_prompt = """You are an assistant for editing Jordanian employment contracts.
You have the current contract and the user's edit request.
You must:
1. Apply the requested edit to the contract
2. Preserve the contract structure and formatting
3. Ensure compliance with Jordanian Labour Law
4. Keep the legal disclaimer at the end
5. Reply with ONLY the complete updated contract, no extra commentary"""

    messages = [  # type: ignore
        {"role": "system", "content": system_prompt},
    ]

    # Add recent chat history for context
    for msg in store.chat_history[-10:]:  # type: ignore
        messages.append({"role": msg["role"], "content": msg["content"]})  # type: ignore

    messages.append({  # type: ignore
        "role": "user",
        "content": f"Current contract:\n{store.current}\n\nEdit request: {user_message}"
    })

    resp = _openai_client.chat.completions.create(
        model="gpt-4o-mini",
        messages=messages,  # type: ignore
        temperature=0,
    )
    updated = (resp.choices[0].message.content or "").strip()
    updated = _ensure_disclaimer(updated, lang)

    # Save version before updating
    store.save_version(meta={"action": "chat_edit", "request": user_message[:100]})
    store.current = updated

    # Parse what changed for the reply
    if lang == "ar":
        reply = f"✅ تم تطبيق التعديل: {user_message[:80]}..."
    else:
        reply = f"✅ Applied edit: {user_message[:80]}..."

    store.add_chat_message("assistant", reply)

    # Re-validate
    issues = validate_contract_data(store.current_data)
    return store.current, reply, issues


# =============================================
# WORD EXPORT (PROFESSIONAL)
# =============================================
def export_to_word(contract_text: str, lang: str = "en") -> str:
    if not contract_text:
        return "❌ No contract to export."

    contract_text = _ensure_disclaimer(contract_text, lang)

    doc = Document()

    # --- Page margins ---
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    lines = contract_text.split("\n")

    for i, line in enumerate(lines):
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue

        # Title (first non-empty line)
        if i == 0 or (i < 3 and line.isupper()):
            p = doc.add_heading(line, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0] if p.runs else p.add_run(line)
            run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

        # Section headers (ALL CAPS or ends with :)
        elif (line.isupper() and len(line) > 3) or (line.endswith(":") and len(line) < 40):
            p = doc.add_heading(line, level=2)
            run = p.runs[0] if p.runs else p.add_run(line)
            run.font.color.rgb = RGBColor(0x16, 0x21, 0x3e)

        # Disclaimer section
        elif "LEGAL DISCLAIMER" in line or "تنويه قانوني" in line:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
            run.font.color.rgb = RGBColor(0xc0, 0x39, 0x2b)
            run.font.size = Pt(10)

        # Disclaimer body
        elif any(w in contract_text[:contract_text.find(line) + 1]
                 for w in ["LEGAL DISCLAIMER", "تنويه قانوني"]) and \
             "LEGAL DISCLAIMER" not in line and "تنويه قانوني" not in line:
            p = doc.add_paragraph(line)
            p.runs[0].font.size = Pt(9)
            p.runs[0].font.color.rgb = RGBColor(0x7f, 0x8c, 0x8d)

        # Regular content
        else:
            p = doc.add_paragraph(line)
            p.runs[0].font.size = Pt(11) if p.runs else Pt(11)

    # Footer
    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.text = f"Generated by Jordanian Legal Assistant | {datetime.now().strftime('%d/%m/%Y')} | Non-binding draft"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer_para.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x95, 0xa5, 0xa6)

    path = os.path.join(EXPORT_DIR, f"contract_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
    doc.save(path)
    return path


# =============================================
# LEGACY COMPATIBILITY (keep old functions working)
# =============================================
def generate_contract(store: ContractStore, data: Dict[str, str]) -> str:
    """Legacy wrapper — calls AI generator, returns contract text only."""
    contract, _ = generate_contract_ai(store, data, lang="en")
    return contract


SECTION_PATTERNS = {
    "salary": r"(Salary:\s*)(.*)",
    "location": r"(Location:\s*)(.*)",
    "job_title": r"(Job Title:\s*)(.*)",
    "termination": r"(TERMINATION:\n)([\s\S]*?)(\n\nLEGAL DISCLAIMER:)",
}


def update_section(store: ContractStore, section: str, new_value: str) -> str:
    """Legacy section updater — still works for direct edits."""
    if not store.current:
        return "❌ No contract exists. Generate a contract first."

    section = (section or "").strip().lower()
    if section not in SECTION_PATTERNS:
        return "❌ Unknown section. Use: salary, location, job_title, termination"

    new_value = (new_value or "").strip()
    store.save_version(meta={"action": "update", "section": section})

    text = store.current

    if section == "termination":
        m = re.search(SECTION_PATTERNS["termination"], text)
        if not m:
            if "LEGAL DISCLAIMER:" in text:
                head, tail = text.split("LEGAL DISCLAIMER:", 1)
                text = head.rstrip() + "\n\nTERMINATION:\n" + new_value + "\n\nLEGAL DISCLAIMER:" + tail
            else:
                text = text.rstrip() + "\n\nTERMINATION:\n" + new_value + "\n\n" + DISCLAIMER_EN
        else:
            text = re.sub(SECTION_PATTERNS["termination"], r"\1" + new_value + r"\3", text)
    else:
        text = re.sub(
            SECTION_PATTERNS[section],
            lambda mm: mm.group(1) + new_value,
            text,
            count=1
        )

    store.current = _ensure_disclaimer(text, store.lang)
    return store.current


# =============================================
# DOCUMENT UPLOAD ANALYSIS
# =============================================
_EXTRACT_FIELDS_PROMPT = """Extract employment contract fields from the text below.
Return ONLY a JSON object with these exact keys (use "" if not found):
employer, employee, job_title, location, contract_type (must be "Limited" or "Unlimited"),
start_date (DD/MM/YYYY), end_date (DD/MM/YYYY or ""), salary (number only, no currency),
hours (number only), days (number only), probation_days (number only or ""),
termination (brief description), benefits (brief description).

Contract text:
{text}

Return valid JSON only, no markdown or explanation."""


def _extract_fields_regex(text: str) -> Dict[str, str]:
    """Regex-based field extractor used as fallback when API is unavailable."""
    fields: Dict[str, str] = {
        k: "" for k in [
            "employer", "employee", "job_title", "location", "contract_type",
            "start_date", "end_date", "salary", "hours", "days",
            "probation_days", "termination", "benefits",
        ]
    }
    salary_m = re.search(r"(?:salary|wage)[:\s]+(\d[\d,\.]+)", text, re.I)
    if salary_m:
        fields["salary"] = re.sub(r"[^\d.]", "", salary_m.group(1))

    hours_m = re.search(r"(\d+)\s*hours?\s*(?:per|a|/)\s*day", text, re.I)
    if hours_m:
        fields["hours"] = hours_m.group(1)

    days_m = re.search(r"(\d+)\s*days?\s*(?:per|a|/)\s*week", text, re.I)
    if days_m:
        fields["days"] = days_m.group(1)

    prob_m = re.search(r"probation[^\d]*(\d+)\s*days?", text, re.I)
    if prob_m:
        fields["probation_days"] = prob_m.group(1)

    if re.search(r"limited|fixed.term", text, re.I):
        fields["contract_type"] = "Limited"
    elif re.search(r"unlimited|indefinite|open.ended", text, re.I):
        fields["contract_type"] = "Unlimited"

    notice_m = re.search(r"(\d+)\s*days?\s*(?:written\s*)?notice", text, re.I)
    if notice_m:
        fields["termination"] = f"{notice_m.group(1)} days written notice"

    return fields


def extract_fields_from_text(text: str) -> Dict[str, str]:
    """
    Extract structured contract fields from raw text.
    Uses GPT-4o-mini for best accuracy; falls back to regex if API unavailable.

    Returns a dict with keys:
      employer, employee, job_title, location, contract_type, start_date,
      end_date, salary, hours, days, probation_days, termination, benefits
    """
    if _openai_client is None or not os.getenv("OPENAI_API_KEY"):
        return _extract_fields_regex(text)

    prompt = _EXTRACT_FIELDS_PROMPT.format(text=text[:6000])
    try:
        import json as _json
        resp = _openai_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            temperature=0,
            response_format={"type": "json_object"},  # type: ignore
        )
        raw = (resp.choices[0].message.content or "{}").strip()
        fields = _json.loads(raw)
        required_keys = [
            "employer", "employee", "job_title", "location", "contract_type",
            "start_date", "end_date", "salary", "hours", "days",
            "probation_days", "termination", "benefits",
        ]
        return {k: str(fields.get(k, "")) for k in required_keys}
    except Exception:
        return _extract_fields_regex(text)


def analyze_uploaded_document(raw_text: str) -> Dict[str, Any]:
    """
    Full analysis pipeline for an uploaded contract document.

    Steps:
      1. Extract structured fields via GPT (or regex fallback)
      2. Run validate_contract_data on extracted fields
      3. Ask GPT to write a brief narrative analysis (optional, requires API key)

    Returns:
      {
        "fields_extracted": dict,
        "validation_issues": List[ValidationIssue],
        "validation_summary": dict,
        "ai_analysis": str,
        "success": bool,
      }
    """
    if not raw_text or not raw_text.strip():
        return {
            "fields_extracted": {},
            "validation_issues": [],
            "validation_summary": get_validation_summary([]),
            "ai_analysis": "No text could be extracted from the uploaded document.",
            "success": False,
        }

    # Step 1: Field extraction
    fields = extract_fields_from_text(raw_text)

    # Step 2: Validation
    issues = validate_contract_data(fields)
    summary = get_validation_summary(issues)

    # Step 3: GPT narrative analysis
    ai_analysis = ""
    if _openai_client and os.getenv("OPENAI_API_KEY"):
        error_lines = "\n".join(
            f"- [{i.severity.upper()}] {i.message}" + (f" ({i.article_ref})" if i.article_ref else "")
            for i in issues
            if i.severity in ("error", "warning")
        ) or "No major issues found."

        analysis_prompt = (
            f"You are a Jordanian Labour Law expert. Analyze this employment contract "
            f"and the issues found. Provide a concise (3-5 sentence) professional summary "
            f"of the contract's compliance status and key risks.\n\n"
            f"Extracted fields:\n{_format_data_for_prompt(fields)}\n\n"
            f"Issues found:\n{error_lines}\n\n"
            f"Write in the same language as the contract (Arabic or English)."
        )
        try:
            resp = _openai_client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": analysis_prompt}],
                temperature=0,
            )
            ai_analysis = (resp.choices[0].message.content or "").strip()
        except Exception:
            ai_analysis = "AI narrative analysis unavailable."
    else:
        ai_analysis = "Set OPENAI_API_KEY for AI-powered narrative analysis."

    return {
        "fields_extracted": fields,
        "validation_issues": issues,
        "validation_summary": summary,
        "ai_analysis": ai_analysis,
        "success": True,
    }
